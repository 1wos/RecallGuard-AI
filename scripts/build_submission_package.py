from __future__ import annotations

import json
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
FINAL = ROOT / "final"
ASSETS = FINAL / "video_assets"
RENDER = FINAL / "rendered_report"

FINAL.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)
RENDER.mkdir(exist_ok=True)

REPORT_DOCX = FINAL / "RecallGuard_AI_Build_Report_Somi.docx"
REPORT_PDF = FINAL / "RecallGuard_AI_Build_Report_Somi.pdf"
VIDEO_MP4 = FINAL / "RecallGuard_AI_Demo_Somi.mp4"
SCRIPT_MD = FINAL / "RecallGuard_AI_Demo_Script_Somi.md"


COLORS = {
    "ink": RGBColor(11, 37, 69),
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "muted": RGBColor(85, 85, 85),
    "light_fill": "F2F4F7",
    "blue_fill": "E8EEF5",
    "green_fill": "EAF4EA",
    "yellow_fill": "FFF4CC",
    "red_fill": "FDEAEA",
    "border": "DADCE0",
}


def load_json(path: str) -> dict:
    return json.loads((ROOT / path).read_text())


def response_text(path: str) -> str:
    data = load_json(path)
    parts = []
    for item in data.get("output", []):
        if item.get("type") == "message":
            for content in item.get("content", []):
                if content.get("type") == "output_text":
                    parts.append(content.get("text", ""))
    return "\n".join(parts).strip()


def extract_json_block(text: str) -> dict | list | None:
    match = re.search(r"```json\s*(.*?)```", text, flags=re.S)
    if not match:
        return None
    try:
        return json.loads(match.group(1))
    except json.JSONDecodeError:
        return None


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_paragraph(p, before=0, after=6, line=1.10, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def run(p, text, size=11, color=None, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(p, before=14 if level == 1 else 10, after=6)
    color = COLORS["blue"] if level <= 2 else COLORS["dark_blue"]
    size = 16 if level == 1 else 13 if level == 2 else 12
    run(p, text, size=size, color=color, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, after=6, line=1.10)
    run(p, text, size=11, color=RGBColor(0, 0, 0))
    return p


def add_callout(doc, title, body, fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    set_table_borders(table, color="B8C7DA")
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    style_paragraph(p, after=4)
    run(p, title, size=11.5, color=COLORS["ink"], bold=True)
    p2 = cell.add_paragraph()
    style_paragraph(p2, after=0, line=1.12)
    run(p2, body, size=10.5, color=RGBColor(0, 0, 0))
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None, header_fill="F2F4F7"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_fill(cell, header_fill)
        set_cell_margins(cell, top=100, bottom=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, after=0)
        run(p, text, size=9.5, color=COLORS["ink"], bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell, top=90, bottom=90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            run(p, str(value), size=9.3, color=RGBColor(0, 0, 0))

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    style_paragraph(header, after=0)
    run(header, "RecallGuard AI | Microsoft Foundry Final Activity", size=9, color=COLORS["muted"])

    footer = section.footer.paragraphs[0]
    style_paragraph(footer, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run(footer, "Somi | Governed Multi-Agent Workflow", size=9, color=COLORS["muted"])
    return doc


def build_report():
    doc = setup_document()

    p = doc.add_paragraph()
    style_paragraph(p, before=18, after=2)
    run(p, "RecallGuard AI", size=26, color=COLORS["ink"], bold=True)
    p = doc.add_paragraph()
    style_paragraph(p, after=12)
    run(p, "Governed Multi-Agent Product Safety Compliance Checker", size=14, color=COLORS["muted"], italic=True)

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Activity", "Build a Governed Multi-Agent Workflow with Microsoft Foundry"],
            ["Scenario", "Marketplace/procurement product safety review"],
            ["Core agents", "Knowledge Agent + Task Agent + Sequential Workflow"],
            ["Foundry project", "recallguard-ai"],
            ["Final task agent", "recallguard-task-agent-v5"],
            ["Submission", "Video MP4 + PDF/DOCX build report"],
        ],
        widths=[1.8, 4.5],
        header_fill=COLORS["blue_fill"],
    )

    add_callout(
        doc,
        "Executive Summary",
        "RecallGuard AI is a governed multi-agent workflow built in Microsoft Foundry. It includes a Knowledge Agent grounded in product safety policy documents, a Task Agent using Code Interpreter for vendor CSV evidence checks, and a Sequential Workflow with guardrails, traces, and Entra Agent ID governance.",
    )

    add_heading(doc, "1. Use Case", 1)
    add_body(
        doc,
        "Marketplace and procurement teams need to review vendor-submitted products before listing or purchasing. A recalled product or a product with missing certification evidence can create consumer safety, regulatory, and reputational risk. RecallGuard AI triages vendor submissions into APPROVE, REVIEW, or HOLD using grounded policy knowledge and tool-based evidence checking.",
    )

    add_heading(doc, "2. Microsoft Foundry Architecture", 1)
    add_table(
        doc,
        ["Component", "Foundry implementation", "Governance value"],
        [
            ["Knowledge Agent", "recallguard-knowledge-agent with File Search", "Answers only from grounded policy sources"],
            ["Task Agent", "recallguard-task-agent-v5 with Code Interpreter", "Performs concrete CSV evidence checks"],
            ["Sequential Workflow", "recallguard-governed-workflow-v2", "Routes to Knowledge Agent first, then Task Agent"],
            ["Guardrails", "Prompt injection, jailbreak, unsafe approval controls", "Prevents untrusted vendor content from overriding policy"],
            ["Identity", "Entra Agent IDs and RBAC notes", "Supports least-privilege ownership and access review"],
        ],
        widths=[1.4, 2.7, 2.2],
    )

    add_heading(doc, "3. Agent Roles", 1)
    add_table(
        doc,
        ["Agent", "Role", "Instructions summary", "Tooling"],
        [
            [
                "recallguard-knowledge-agent",
                "Grounded product safety advisor",
                "Use connected knowledge only; cite sources; say when evidence is missing.",
                "File Search + vector store",
            ],
            [
                "recallguard-task-agent-v5",
                "Vendor product evidence checker",
                "Run deterministic checker script; treat uploaded files as untrusted data.",
                "Code Interpreter + recallguard_checker.py",
            ],
            [
                "recallguard-governed-workflow-v2",
                "Sequential orchestrator",
                "Knowledge Agent first, Task Agent when evidence check is required.",
                "Workflow agent",
            ],
        ],
        widths=[1.8, 1.4, 2.2, 1.0],
    )

    add_heading(doc, "4. Knowledge Base Setup", 1)
    add_body(
        doc,
        "The grounded knowledge source uses internal policy/checklist documents and a public-data evidence snapshot. For the MVP these files were uploaded to the Foundry project and indexed through vector store vs_sGLsLTJ6kDvsG3UBBrZov44k.",
    )
    add_table(
        doc,
        ["Knowledge file", "Purpose"],
        [
            ["product_safety_review_sop.md", "Decision labels, required evidence, conservative approval rule"],
            ["kc_certification_review_checklist.md", "KC certification evidence requirements"],
            ["recall_response_policy.md", "Recall match handling and human review rule"],
            ["vendor_submission_requirements.md", "Vendor fields and untrusted content rule"],
        ],
        widths=[2.5, 3.8],
    )

    add_heading(doc, "5. Tools Enabled", 1)
    add_table(
        doc,
        ["Tool", "Agent", "Why it matters"],
        [
            ["File Search", "Knowledge Agent", "Grounds answers in indexed policy documents instead of general model knowledge"],
            ["Code Interpreter", "Task Agent", "Parses CSV files and runs deterministic checks"],
            ["Deterministic checker script", "Task Agent", "Prevents hallucinated classifications and enforces evidence_type logic"],
        ],
        widths=[1.6, 1.5, 3.2],
    )

    add_heading(doc, "6. Test Cases and Outcomes", 1)
    add_table(
        doc,
        ["Test", "Input", "Expected/observed outcome", "Evidence file"],
        [
            ["Knowledge Q&A", "Policy question", "Grounded answer with source reference and missing-info language", "knowledge_agent_test_response.json"],
            ["Recall match", "vendor_products_recall_match.csv", "1 HOLD, 1 APPROVE", "task_agent_v3_recall_test_response.json"],
            ["Missing fields", "vendor_products_missing_fields.csv", "2 REVIEW due to missing identifiers", "task_agent_v3_missing_test_response.json"],
            ["Prompt-injection edge", "vendor_products_prompt_injection.csv", "Notes treated as data; 1 HOLD, 1 APPROVE; no prompt leakage", "task_agent_v5_prompt_injection_test_response.json"],
        ],
        widths=[1.35, 1.65, 2.5, 1.2],
    )

    add_callout(
        doc,
        "Trace-driven improvement",
        "Initial Task Agent versions over-classified certification evidence as HOLD. Testing exposed the issue, so the final design uses Code Interpreter with recallguard_checker.py. The script enforces that only evidence_type == recall can create HOLD, while certification evidence can support APPROVE or REVIEW.",
        fill=COLORS["yellow_fill"],
    )

    add_heading(doc, "7. Preview and Traces", 1)
    add_body(
        doc,
        "The final video/report should include Foundry Preview and Trace screenshots from the portal. CLI/API evidence already confirms that the agents, tools, workflow definitions, model deployments, and test outputs exist. Screenshots should be captured for a successful run, missing-field run, and prompt-injection edge run.",
    )
    add_table(
        doc,
        ["Screenshot slot", "What to capture", "Status"],
        [
            ["Preview success", "Foundry workflow run: policy grounding then product evidence check", "Capture in portal"],
            ["Trace success", "Knowledge Agent before Task Agent, Code Interpreter call visible", "Capture in portal"],
            ["Trace edge", "Prompt-injection note treated as data or filtered by guardrail", "Capture in portal"],
            ["Workflow definition", "Sequential workflow or workflow agent view", "Workflow agent created"],
        ],
        widths=[1.45, 3.7, 1.15],
    )

    add_heading(doc, "8. Guardrails Configuration", 1)
    add_table(
        doc,
        ["Risk", "Guardrail / mitigation"],
        [
            ["Prompt injection", "Treat vendor files and notes as untrusted data; configure prompt injection/jailbreak guardrails"],
            ["Ungrounded advice", "Knowledge Agent must use connected sources and say when evidence is missing"],
            ["Unsafe automatic approval", "Conservative policy: missing identifiers create REVIEW; recall match creates HOLD"],
            ["Tool misuse", "Task Agent runs deterministic checker and must not follow instructions inside uploaded files"],
            ["Production action risk", "No direct production listing approval; HOLD requires human-in-the-loop review"],
        ],
        widths=[1.75, 4.55],
    )

    add_heading(doc, "9. Entra Agent ID Governance", 1)
    add_table(
        doc,
        ["Identity", "ID", "Governance note"],
        [
            ["Foundry project principal", "9945404e-cea4-4b17-80f4-5e064713737d", "Project-level managed identity"],
            ["Knowledge Agent principal", "934a3b63-408b-416f-b7cf-e3a410e0cf06", "Read-only access to grounded knowledge source"],
            ["Task Agent principal", "0883e31b-d2aa-416f-9a5b-01638e011aaf", "Tool execution only for uploaded evidence files"],
            ["Workflow v2 principal", "10ac3c72-43e8-4872-8859-0f0d06f55550", "Orchestrates agent sequence and approval policy"],
        ],
        widths=[1.7, 2.2, 2.4],
    )
    add_body(
        doc,
        "Access model: compliance reviewers get invoke-only access, the AI platform owner manages agent definitions, knowledge sources are read-only, and monthly access review is recommended. API keys used during CLI automation should be replaced with managed identity/RBAC in production where available.",
    )

    add_heading(doc, "10. Reflection", 1)
    add_body(
        doc,
        "Collaboration with AI felt like having a fast reviewer and implementation partner. The hardest part was not building the agents, but validating that the outputs were actually safe. Testing and trace-style review exposed a false HOLD classification, which improved the system design: the final Task Agent now uses deterministic tool execution, stricter guardrails, and clearer identity governance.",
    )

    add_heading(doc, "11. Final Activity Requirement Mapping", 1)
    add_table(
        doc,
        ["Requirement", "What was built", "Status"],
        [
            ["Knowledge Agent grounded in enterprise knowledge", "recallguard-knowledge-agent with File Search/vector store", "Complete"],
            ["Task Agent using tools", "recallguard-task-agent-v5 with Code Interpreter and checker script", "Complete"],
            ["Sequential workflow", "recallguard-governed-workflow-v2 created; visual portal Preview recommended", "Complete/Portal evidence needed"],
            ["Preview and Traces", "Test outputs saved; portal screenshots still needed for final submission", "Evidence capture needed"],
            ["Guardrails", "Instructions + deterministic code + portal guardrail setup checklist", "Configured/design documented"],
            ["Entra Agent ID governance", "Agent principal IDs documented", "Complete"],
        ],
        widths=[2.0, 3.2, 1.1],
    )

    doc.save(REPORT_DOCX)


def convert_docx_to_pdf():
    cmd = [
        "/opt/homebrew/bin/soffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(FINAL),
        str(REPORT_DOCX),
    ]
    subprocess.run(cmd, check=True, cwd=str(ROOT))


def draw_wrapped(draw, xy, text, font, fill, max_width, line_spacing=8):
    x, y = xy
    words = text.split()
    lines = []
    line = ""
    for word in words:
        test = f"{line} {word}".strip()
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] <= max_width or not line:
            line = test
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_spacing
    return y


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for path in candidates:
        if path and Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def make_slide(index, title, subtitle, bullets, footer):
    img = Image.new("RGB", (1920, 1080), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1920, 120], fill="#0B2545")
    draw.text((90, 38), "RecallGuard AI", font=font(42, True), fill="#FFFFFF")
    draw.text((1510, 46), "Microsoft Foundry", font=font(30, False), fill="#D9EAF7")
    draw.rectangle([80, 180, 1840, 935], fill="#FFFFFF", outline="#DADCE0", width=3)
    draw.text((120, 220), title, font=font(54, True), fill="#0B2545")
    y = draw_wrapped(draw, (122, 298), subtitle, font(31), "#475569", 1600, 10)
    y += 42
    for bullet in bullets:
        draw.ellipse([125, y + 10, 145, y + 30], fill="#2E74B5")
        y = draw_wrapped(draw, (165, y), bullet, font(34), "#111827", 1540, 12)
        y += 24
    draw.text((90, 990), footer, font=font(24), fill="#64748B")
    path = ASSETS / f"slide_{index:02d}.png"
    img.save(path)
    return path


def build_video():
    slides = [
        (
            "Use case",
            "RecallGuard AI reviews vendor product submissions before marketplace listing or procurement approval.",
            [
                "Built in Microsoft Foundry as a governed multi-agent workflow.",
                "Uses Korea product safety recall and certification evidence.",
                "Classifies products as APPROVE, REVIEW, or HOLD.",
            ],
            "0:00 - Scenario and goal",
            "This is RecallGuard AI, a governed multi-agent product safety compliance checker built in Microsoft Foundry. It helps marketplace and procurement teams review vendor product submissions against safety policy and recall evidence.",
        ),
        (
            "Knowledge Agent",
            "The Knowledge Agent answers only from grounded policy documents and says when evidence is missing.",
            [
                "Agent: recallguard-knowledge-agent.",
                "Knowledge source: product safety SOP, KC checklist, recall response policy, vendor requirements.",
                "Tool: File Search with vector store.",
            ],
            "0:20 - Grounded knowledge",
            "The Knowledge Agent answers only from grounded knowledge sources. It uses File Search over the product safety SOP, KC certification checklist, recall response policy, and vendor submission requirements.",
        ),
        (
            "Task Agent",
            "The Task Agent performs real work with Code Interpreter instead of guessing from the prompt.",
            [
                "Agent: recallguard-task-agent-v5.",
                "Tool: Code Interpreter.",
                "Deterministic checker: recallguard_checker.py.",
            ],
            "0:50 - Tool-based action",
            "The Task Agent performs the concrete action. It uses Code Interpreter and a deterministic checker script to inspect vendor CSV files and compare them with recall and certification evidence.",
        ),
        (
            "Sequential Workflow",
            "The workflow routes the request to the Knowledge Agent first, then the Task Agent when file checking is required.",
            [
                "Workflow Agent: recallguard-governed-workflow-v2.",
                "Pattern: sequential orchestration.",
                "HOLD decisions require human compliance review.",
            ],
            "1:20 - Orchestration",
            "The sequential workflow grounds the review policy first, then invokes the Task Agent when evidence checking is required. Any HOLD decision requires human compliance review before listing.",
        ),
        (
            "Successful and failure tests",
            "The final Task Agent was tested with recall, missing-field, and prompt-injection edge cases.",
            [
                "Recall match: one HOLD and one APPROVE.",
                "Missing fields: two REVIEW decisions.",
                "Prompt-injection edge: notes treated only as data.",
            ],
            "1:50 - Tests and outcomes",
            "I tested a recall match case, a missing-field case, and a prompt-injection edge case. The final system produced one hold and one approval for the recall batch, two reviews for missing identifiers, and safe handling of vendor notes as data.",
        ),
        (
            "Trace-driven improvement",
            "Testing exposed a false HOLD classification, so the Task Agent was hardened with deterministic code.",
            [
                "Early versions over-classified certification evidence as HOLD.",
                "Fix: only evidence_type equals recall can create HOLD.",
                "Certification evidence can support APPROVE or REVIEW only.",
            ],
            "2:15 - What improved after testing",
            "The hardest part was validating that agent outputs were actually safe. Testing exposed a false hold classification, so I improved the design with deterministic code and stricter evidence handling.",
        ),
        (
            "Guardrails and identity",
            "The design includes prompt-injection protection, untrusted-file handling, and Entra Agent ID governance.",
            [
                "Guardrails: prompt injection, jailbreak, ungrounded claims, unsafe approvals.",
                "Identity: Entra Agent IDs documented for project, agents, and workflow.",
                "Access model: least privilege and invoke-only reviewers.",
            ],
            "2:35 - Governance",
            "The governance layer includes prompt-injection protection, untrusted-file handling, Entra Agent ID notes, RBAC ownership, and least-privilege access. The agent does not directly approve production listings.",
        ),
        (
            "Reflection",
            "AI collaboration felt fast, but traces and tests were essential to make the workflow trustworthy.",
            [
                "AI accelerated ideation, implementation, and review.",
                "The hard part was safety validation and evidence discipline.",
                "The final design is more governed, traceable, and tool-grounded.",
            ],
            "2:50 - Reflection",
            "Collaboration with AI felt like having a fast implementation partner. But the most valuable part was review: traces and tests helped reveal failure modes and improve the workflow with guardrails, deterministic tools, and identity governance.",
        ),
    ]

    script_lines = ["# RecallGuard AI Demo Script\n"]
    segment_paths = []
    for i, (title, subtitle, bullets, footer, narration) in enumerate(slides, start=1):
        slide = make_slide(i, title, subtitle, bullets, footer)
        script_lines.append(f"## Slide {i}: {title}\n\n{narration}\n")
        audio = ASSETS / f"slide_{i:02d}.aiff"
        subprocess.run(["say", "-v", "Samantha", "-r", "178", "-o", str(audio), narration], check=True)
        segment = ASSETS / f"segment_{i:02d}.mp4"
        subprocess.run(
            [
                "ffmpeg",
                "-y",
                "-loop",
                "1",
                "-i",
                str(slide),
                "-i",
                str(audio),
                "-shortest",
                "-c:v",
                "libx264",
                "-tune",
                "stillimage",
                "-pix_fmt",
                "yuv420p",
                "-c:a",
                "aac",
                "-b:a",
                "128k",
                str(segment),
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        segment_paths.append(segment)

    SCRIPT_MD.write_text("\n".join(script_lines))

    concat_file = ASSETS / "segments.txt"
    concat_file.write_text("".join(f"file '{p.name}'\n" for p in segment_paths))
    subprocess.run(
        [
            "ffmpeg",
            "-y",
            "-f",
            "concat",
            "-safe",
            "0",
            "-i",
            str(concat_file.name),
            "-c",
            "copy",
            str(VIDEO_MP4),
        ],
        cwd=str(ASSETS),
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )


def main():
    build_report()
    convert_docx_to_pdf()
    build_video()
    print(json.dumps({
        "docx": str(REPORT_DOCX),
        "pdf": str(REPORT_PDF),
        "video": str(VIDEO_MP4),
        "script": str(SCRIPT_MD),
    }, indent=2))


if __name__ == "__main__":
    main()
